import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

DIRS = [
    "1-SASI",
    "2-Politicas_Seguridad",
    "3-Controles_Acceso",
    "4-Continuidad_Negocio",
    "5-Gestion_Riesgos",
    "6-Control_Operaciones",
    "portal-tic/public"
]

def add_border(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'<w:tcBorders %s>'
                                  r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                                  r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                                  r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                                  r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                                  r'</w:tcBorders>' % nsdecls('w'))
            tcPr.append(tcBorders)

def generate_markdown(filename, data):
    md = f"![Logo BNPHU](/logo-bnphu.png)\n\n# BIBLIOTECA NACIONAL PEDRO HENRÍQUEZ UREÑA\n**DPyD | Versión 1**\n**MANUAL DE POLÍTICAS Y PROCEDIMIENTOS**\n\n---\n\n"
    md += f"**NOMBRE DEL PROCESO:** {data['nombre']}\n\n"
    md += f"| Preparado por: {data['preparado']} | Aprobado por: {data['aprobado']} | Código/Páginas: {data.get('codigo', 'BNPHU-TIC-00X')} |\n|---|---|---|\n\n"
    md += f"**1.0 Propósito o Misión:**\n{data['proposito']}\n\n"
    md += f"**2.0 Alcance:**\n- **Empieza:** {data['alcance_empieza']}\n- **Incluye:** {data['alcance_incluye']}\n- **Termina:** {data['alcance_termina']}\n\n"
    
    md += "**3.0 Dueño o responsables:**\n"
    for r in data['responsables']:
        md += f"{r}\n"
    md += "\n"
    
    md += "**4.0 Documentos de referencia:**\n"
    for ref in data['referencias']:
        md += f"{ref}\n"
    md += "\n"
    
    md += "**5.0 Políticas del Procedimiento:**\n"
    for p in data['politicas']:
        md += f"{p}\n"
    md += "\n"
    
    md += "**6.0 Descripción de las Actividades del Proceso:**\n\n"
    md += "| Responsable | Descripción |\n|---|---|\n"
    for r, d in data['actividades']:
        md += f"| **{r}** | {d} |\n"
    md += "\n"
    
    md += "**7.0 ANEXOS:**\n"
    for anexo in data['anexos']:
        md += f"- {anexo}\n"
    md += "\n"
    
    md += "**8.0 REGISTROS:**\n| CODIGO | NOMBRE | ALMACENAMIENTO | ARCHIVADO | TIEMPO | DISPOSICION |\n|---|---|---|---|---|---|\n"
    md += "| SEC-DOC | Registro Dpto TIC | Servidor Seguro | Digital/Impreso | Permanente | Restringido |\n\n"
    
    md += "**9.0 Historia de Cambio:**\n| REVISIONES | FECHAS | SECCION | DESCRIPCION | REVISADO POR | REFRENDADO POR |\n|---|---|---|---|---|---|\n"
    md += "| 1.0 | Actual | Todas | Creación oficial | Dpto. TIC | Dirección |\n\n"
    
    md += f"**10.0 TIEMPO DE RESPUESTA:** {data['tiempo']}\n\n"
    md += "**FIN DEL PROCEDIMIENTO**\n"
    
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(md)

def create_procedure(filename_docx, filename_md, data):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    header_table = doc.add_table(rows=2, cols=3)
    header_table.style = 'Table Grid'
    
    a = header_table.cell(0, 0)
    b = header_table.cell(1, 0)
    a.merge(b)
    a.text = ""
    p = a.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if os.path.exists("logo-bnphu.png"):
        run.add_picture("logo-bnphu.png", width=Inches(1.2))
    else:
        run.add_text("BNPHU")
    
    c = header_table.cell(0, 1)
    c.text = "BIBLIOTECA NACIONAL PEDRO HENRÍQUEZ UREÑA"
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    d = header_table.cell(1, 1)
    d.text = "MANUAL DE POLÍTICAS Y PROCEDIMIENTOS"
    d.paragraphs[0].runs[0].bold = True
    d.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    e = header_table.cell(0, 2)
    f = header_table.cell(1, 2)
    e.merge(f)
    e.text = "DPyD\nVersión 1:\nEnero, 2020"
    e.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set explicit column widths for header table
    header_table.autofit = False
    header_table.allow_autofit = False
    for r in header_table.rows:
        r.cells[0].width = Inches(1.5)
        r.cells[1].width = Inches(4.5)
        r.cells[2].width = Inches(1.5)

    doc.add_paragraph() 
    
    main_table = doc.add_table(rows=0, cols=3)
    main_table.autofit = False
    main_table.allow_autofit = False
    
    def set_main_row_widths(r):
        r.cells[0].width = Inches(2.5)
        r.cells[1].width = Inches(3.0)
        r.cells[2].width = Inches(2.0)

    main_table.style = 'Table Grid'

    row = main_table.add_row()
    row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    row.cells[0].text = f"NOMBRE DEL PROCESO: {data['nombre']}"
    row.cells[0].paragraphs[0].runs[0].bold = True

    row = main_table.add_row()
    set_main_row_widths(row)
    row.cells[0].text = f"Preparado por: {data['preparado']}"
    row.cells[1].text = f"Aprobado por: {data['aprobado']}"
    row.cells[2].text = f"Código/Páginas:\n{data.get('codigo', 'BNPHU-TIC-00X')}"

    row = main_table.add_row()
    set_main_row_widths(row)
    row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    row.cells[0].text = f"1.0 Propósito o Misión:\n{data['proposito']}"

    row = main_table.add_row()
    row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    row.cells[0].text = f"2.0 Alcance:\nEmpieza: {data['alcance_empieza']}\nIncluye: {data['alcance_incluye']}\nTermina: {data['alcance_termina']}"

    row = main_table.add_row()
    row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    row.cells[0].text = "3.0 Dueño o responsables:\n" + "\n".join(data['responsables'])

    row = main_table.add_row()
    row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    row.cells[0].text = "4.0 Documentos de referencia:\n" + "\n".join(data['referencias'])

    row = main_table.add_row()
    row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    row.cells[0].text = "5.0 Políticas del Procedimiento:\n" + "\n".join(data['politicas'])

    row = main_table.add_row()
    row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    p = row.cells[0].paragraphs[0]
    p.add_run("6.0 Descripción de las Actividades del Proceso:\n").bold = True
    
    act_table = row.cells[0].add_table(rows=1, cols=2)
    act_table.style = 'Table Grid'
    act_table.cell(0,0).text = "Responsable"
    act_table.cell(0,1).text = "Descripción"
    for r, d in data['actividades']:
        r_row = act_table.add_row()
        r_row.cells[0].text = r
        r_row.cells[1].text = d

    row = main_table.add_row()
    row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    row.cells[0].text = "7.0 ANEXOS:\n" + "\n".join(data['anexos'])

    row = main_table.add_row()
    row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    p = row.cells[0].paragraphs[0]
    p.add_run("8.0 REGISTROS:\n").bold = True
    reg_table = row.cells[0].add_table(rows=1, cols=6)
    reg_table.style = 'Table Grid'
    headers = ["CODIGO", "NOMBRE", "ALMACENAMIENTO", "ARCHIVADO", "TIEMPO", "DISPOSICION"]
    for i, h in enumerate(headers):
        reg_table.cell(0,i).text = h

    row = main_table.add_row()
    row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    p = row.cells[0].paragraphs[0]
    p.add_run("9.0 Historia de Cambio:\n").bold = True
    hist_table = row.cells[0].add_table(rows=1, cols=6)
    hist_table.style = 'Table Grid'
    headers_hist = ["REVISIONES", "FECHAS", "SECCION", "DESCRIPCION", "REVISADO POR", "REFRENDADO POR"]
    for i, h in enumerate(headers_hist):
        hist_table.cell(0,i).text = h

    row = main_table.add_row()
    row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    row.cells[0].text = f"10.0 TIEMPO DE RESPUESTA: {data['tiempo']}"

    row = main_table.add_row()
    row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    p = row.cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("FIN DEL PROCEDIMIENTO").bold = True

    add_border(header_table)
    add_border(main_table)

    doc.save(filename_docx)
    generate_markdown(filename_md, data)

# ----------------- DATOS DE LOS PROCEDIMIENTOS COMPLETOS -----------------

sasi_data = {
    "nombre": "Sistema para la Administración de la Seguridad de la Información (SASI)",
    "codigo": "BNPHU-TIC-001",
    "preparado": "Departamento de TIC",
    "aprobado": "Dirección Nacional",
    "proposito": "Garantizar la protección de los activos de información de la BNPHU mediante la implementación estructurada y continua de políticas y directrices que aseguren la confidencialidad, integridad y disponibilidad, alineados con el marco NORTIC A7 y el modelo NIST CSF.",
    "alcance_empieza": "Definición oficial del alcance y emisión de las políticas generales de seguridad.",
    "alcance_incluye": "Diseño del marco regulatorio, inventario de activos, análisis de riesgos y declaración de aplicabilidad (SoA).",
    "alcance_termina": "Ejecución de auditorías, revisión por la alta gerencia y planes de mejora continua.",
    "responsables": ["3.1 Máxima Autoridad Ejecutiva (MAE)", "3.2 Oficial en Jefe de Seguridad de Información (CISO)", "3.3 Encargados de Áreas"],
    "referencias": ["4.1 NORTIC A7:2016 (Cap. I y II)", "4.2 NIST CSF v2.0 (Govern: GV.OC, GV.RM)", "4.3 SISTICGE", "4.4 iTICge", "4.5 ISO 27001"],
    "politicas": [
        "5.1 La MAE es la responsable máxima de garantizar los recursos necesarios para el SASI.",
        "5.2 El SASI debe estar alineado a los objetivos estratégicos y organizacionales de la BNPHU, apoyado en el Plan Estratégico Institucional (PEI).",
        "5.3 El organismo mantendrá una Colección de Documentos Legales que rijan las operaciones y acuerdos de confidencialidad institucionales.",
        "5.4 Todo personal deberá recibir capacitaciones periódicas en ciberseguridad, enfocadas en los roles definidos dentro del SASI.",
        "5.5 Se requiere una Declaración de Aplicabilidad que indique cuáles de los controles de la norma aplican a la BNPHU.",
        "5.6 El sistema será evaluado como mínimo una vez al año mediante auditorías internas o externas."
    ],
    "actividades": [
        ("Alta Dirección / MAE", "6.1 Define y firma formalmente la Misión del organismo respecto a la ciberseguridad.\n6.2 Asigna los recursos operativos y designa al CISO."),
        ("CISO / Dpto. TIC", "6.3 Elabora el documento de definición y alcance del SASI identificando todos los sistemas de información.\n6.4 Identifica amenazas, vulnerabilidades, probabilidades e impactos."),
        ("CISO / Dpto. TIC", "6.5 Define e implementa controles para el tratamiento de riesgos y los plasma en la Declaración de Aplicabilidad.\n6.6 Desarrolla métricas y KPIs para medir la efectividad de los controles implementados."),
        ("Dpto. TIC / Recursos Humanos", "6.7 Organiza y documenta capacitaciones de concienciación sobre ciberseguridad a todo el personal."),
        ("CISO / Dpto. TIC", "6.8 Monitorea todas las incidencias y realiza evaluaciones periódicas de riesgos para garantizar la Mejora Continua del SASI.")
    ],
    "anexos": ["7.1 Documento de Alcance del SASI", "7.2 Matriz de Declaración de Aplicabilidad", "7.3 Catálogo de Servicios e Inventario de Activos", "7.4 Actas de Concienciación"],
    "tiempo": "La revisión del SASI se realiza anualmente o ante cambios significativos."
}

admin_info_data = {
    "nombre": "Políticas para la Administración y Tratamiento Seguro de la Información",
    "codigo": "BNPHU-TIC-002",
    "preparado": "Departamento de TIC",
    "aprobado": "Dirección Nacional",
    "proposito": "Establecer las normativas relativas a las responsabilidades de los empleados frente a la información institucional, su correcta clasificación, almacenamiento, retención, respaldo, recuperación y eliminación definitiva.",
    "alcance_empieza": "Creación, recepción o digitalización de la información.",
    "alcance_incluye": "Clasificación de información, definición de cuotas, retención temporal, encriptación, respaldos periódicos y destrucción certificada.",
    "alcance_termina": "Borrado seguro o destrucción física de la información o equipos.",
    "responsables": ["3.1 Encargado de TIC", "3.2 Analista de Recursos Humanos", "3.3 Responsable de Acceso a la Información (RAI)"],
    "referencias": ["4.1 NORTIC A7:2016 (Sec. 2.02 y 2.03)", "4.2 NIST CSF v2.0 (Protect - Data Security PR.DS)", "4.3 Ley 200-04 Libre Acceso a la Info. Pública", "4.4 Ley General de Archivo No. 481-08"],
    "politicas": [
        "5.1 Responsabilidad del Empleado: Prohibido compartir credenciales. La estación de trabajo debe ser bloqueada al ausentarse. No se permite compartir información fuera de los canales oficiales.",
        "5.2 Clasificación de Información: Todo dato se clasifica en 1) Pública, 2) Valiosa, 3) Sensitiva, o 4) Confidencial. La información sensitiva y confidencial debe estar cifrada y restringida a personal autorizado.",
        "5.3 Divulgación: Las solicitudes ciudadanas de información deben canalizarse vía la Oficina de Acceso a la Información (OAI).",
        "5.4 Almacenamiento: Se deben utilizar rutas dedicadas y auditadas para archivos clasificados. Se asignan cuotas por usuario. No almacenar audio o video personal.",
        "5.5 Respaldos (Backups): Deben hacerse copias de seguridad incrementales, diferenciales o completas, mantenidas en sitios fuera de la localidad principal, y probarse su restauración semestralmente.",
        "5.6 Borrado Seguro: La destrucción física requiere trituración o desintegración. La lógica requiere programas certificados de wiping (sobreescritura). Todo descarte debe documentarse."
    ],
    "actividades": [
        ("RRHH / Usuario", "6.1 Firma el documento de entendimiento de políticas al ingresar a la BNPHU.\n6.2 Bloquea su equipo cada vez que se ausenta de su área."),
        ("Dpto. TIC", "6.3 Implementa GPOs para bloqueo automático de sesión tras 15 min de inactividad.\n6.4 Asigna y revisa cuotas de disco en los servidores de almacenamiento."),
        ("OAI / Jurídico", "6.5 Canaliza, evalúa y responde las solicitudes de acceso a la información del ciudadano."),
        ("Dpto. TIC (Base de Datos / Redes)", "6.6 Ejecuta los cronogramas de backups automáticos (diarios, semanales, mensuales).\n6.7 Verifica la integridad realizando pruebas de restauración semestrales."),
        ("Dpto. TIC / Bienes Nacionales", "6.8 Ante un descargo, se ejecuta herramienta de desmagnetización o Wiping. Si es medio físico (papel/CD), se realiza trituración documentada mediante actas.")
    ],
    "anexos": ["7.1 Acuerdo de Compromiso del Empleado", "7.2 Matriz de Clasificación de la Información", "7.3 Actas de Borrado Seguro", "7.4 Calendario de Backups"],
    "tiempo": "Aplicación permanente. Procedimientos de restauración deben tomar máximo 24 horas."
}

acceso_data = {
    "nombre": "Administración y Control de Accesos",
    "codigo": "BNPHU-TIC-003",
    "preparado": "Departamento de TIC",
    "aprobado": "Dirección Nacional",
    "proposito": "Prevenir accesos no autorizados a la información e infraestructuras, mediante la gestión rigurosa de privilegios lógicos y controles físicos basados en el principio de necesidad de saber.",
    "alcance_empieza": "Solicitud de acceso físico a instalaciones críticas o solicitud lógica a sistemas.",
    "alcance_incluye": "Creación, revisión, revocación de usuarios, gestión de contraseñas, configuraciones VPN, protección física de servidores y centros de datos.",
    "alcance_termina": "Cierre o revocación del acceso, con sus respectivas bitácoras de auditoría conservadas.",
    "responsables": ["3.1 Administrador de Sistemas", "3.2 Administrador de Redes", "3.3 Seguridad Física", "3.4 MAE"],
    "referencias": ["4.1 NORTIC A7:2016 (Capítulo III)", "4.2 NIST CSF v2.0 (Protect - Identity Management and Access Control PR.AC)"],
    "politicas": [
        "5.1 Controles de Accesos de Usuarios: Requiere aprobación del superior jerárquico. Todo movimiento (nuevo, cambio, salida) debe ser notificado inmediatamente por RRHH.",
        "5.2 Control Físico de Infraestructura: Niveles de acceso del 1 al 4. El Centro de Datos es Nivel 4 (acceso estrictamente restringido a TIC o autorizados y acompañados). Todo acceso físico a áreas sensibles debe registrarse en bitácora.",
        "5.3 Política de Contraseñas: Mínimo 8 caracteres, alfanuméricas, incluir mayúsculas y minúsculas. Deben cambiarse cada cuarenta y cinco (45) días.",
        "5.4 Conexiones Remotas e Internet: Las VPN deben utilizar cifrado de mínimo 128 bits (IPsec / TLS). Solo personal autorizado utilizará VPN.",
        "5.5 Navegación y Filtros: Prohibido el acceso a sitios ilícitos, juegos, apuestas o pornografía mediante firewalls/proxies corporativos."
    ],
    "actividades": [
        ("RRHH / Dpto. Solicitante", "6.1 Envía formulario de alta o baja de empleado a la mesa de ayuda (TIC) para gestionar la cuenta de Active Directory."),
        ("Dpto. TIC (Administrador de Sistemas)", "6.2 Crea la cuenta asignando los permisos mínimos necesarios en función del rol.\n6.3 Deshabilita la cuenta inmediatamente tras recibir notificación de salida y resguarda los datos por tiempo legal."),
        ("Dpto. TIC / Seguridad Física", "6.4 Mantiene bitácora de firmas y cédulas para cualquier visitante o externo que requiera ingresar al Centro de Datos.\n6.5 Acompaña a todo personal externo durante la visita a zonas de Nivel 4."),
        ("Dpto. TIC (Administrador de Redes)", "6.6 Implementa filtrado web (Proxy/Firewall) para denegar accesos a URL maliciosas o prohibidas.\n6.7 Configura y monitorea conexiones VPN, verificando certificados digitales de acceso.")
    ],
    "anexos": ["7.1 Formulario de Alta/Baja de Usuarios", "7.2 Bitácora Física de Acceso a Centro de Datos", "7.3 Matriz de Perfiles de Acceso Lógico"],
    "tiempo": "Altas y bajas en 24 horas. Auditoría de permisos anualmente."
}

continuidad_data = {
    "nombre": "Plan de Disponibilidad y Continuidad del Negocio (BCP / DRP)",
    "codigo": "BNPHU-TIC-004",
    "preparado": "Departamento de TIC",
    "aprobado": "Dirección Nacional",
    "proposito": "Mantener las operaciones tecnológicas a un nivel aceptable frente a eventualidades o desastres, estableciendo un protocolo de gestión de incidentes y un esquema de recuperación técnica.",
    "alcance_empieza": "Detección de evento, interrupción o alerta tecnológica.",
    "alcance_incluye": "Análisis de capacidad, registro del incidente, activación del Comité de Continuidad (CONTI), Análisis de Impacto (BIA), y recuperación de desastres (DRP).",
    "alcance_termina": "Restauración a la normalidad operativa, elaboración de lecciones aprendidas y cierre del evento.",
    "responsables": ["3.1 Comité de Continuidad (CONTI)", "3.2 Oficial de Seguridad", "3.3 Departamento de TIC"],
    "referencias": ["4.1 NORTIC A7:2016 (Capítulo IV)", "4.2 NIST CSF v2.0 (Respond - Incident Management RS.MA / Recover RC.RP)"],
    "politicas": [
        "5.1 Gestión de Incidentes (PGI): Todo incidente debe reportarse. Incidentes criminales deben escalarse al DICAT (Ley 53-07).",
        "5.2 Comité de Continuidad (CONTI): Integrado por la MAE, Enc. de TIC y áreas claves. Toma las decisiones estratégicas ante desastres.",
        "5.3 BIA (Análisis de Impacto): Identificar los procesos críticos institucionales y establecer el RTO (Recovery Time Objective) y RPO (Recovery Point Objective).",
        "5.4 DRP (Disaster Recovery Plan): Incluye contingencias técnicas, respaldo de infraestructuras críticas, mantenimientos o localidades alternas con al menos 2 semanas de autonomía.",
        "5.5 Simulacros: Es obligatorio documentar y ejecutar simulacros de evacuación, fallo tecnológico y restauración al menos una vez al año."
    ],
    "actividades": [
        ("Usuario / Dpto. TIC", "6.1 Detecta la interrupción y registra un ticket de incidente. TIC inicia el diagnóstico primario."),
        ("Dpto. TIC / CISO", "6.2 Contiene el incidente. Si sobrepasa la capacidad operativa normal, se escala al Comité CONTI."),
        ("Comité CONTI / MAE", "6.3 Analiza el nivel de impacto según el BIA. Declara formalmente estado de emergencia o desastre y activa el DRP."),
        ("Dpto. TIC", "6.4 Ejecuta los procedimientos de contingencia: restauración de respaldos off-site, conmutación a enlaces redundantes o activación de servicios en la nube."),
        ("Comité CONTI", "6.5 Concluida la recuperación, elabora informe de lecciones aprendidas, ajusta el DRP y envía reportes formales a organismos rectores (OGTIC, CNCS).")
    ],
    "anexos": ["7.1 Plan de Recuperación ante Desastres (DRP)", "7.2 Documento de BIA (Análisis de Impacto)", "7.3 Informes de Simulacros"],
    "tiempo": "Ejecución inmediata. Tiempos de recuperación varían según el RTO del proceso."
}

riesgos_data = {
    "nombre": "Metodología y Gestión de Riesgos de Ciberseguridad",
    "codigo": "BNPHU-TIC-005",
    "preparado": "Departamento de TIC",
    "aprobado": "Dirección Nacional",
    "proposito": "Definir un enfoque estructurado para identificar, analizar, evaluar y tratar los riesgos tecnológicos, garantizando que estén bajo umbrales aceptables para la BNPHU.",
    "alcance_empieza": "Inventariado de activos y definición de la metodología de evaluación.",
    "alcance_incluye": "Identificación de amenazas, cálculo de probabilidad e impacto, determinación de riesgo inherente, aplicación de controles y monitoreo del riesgo residual.",
    "alcance_termina": "Documento de aceptación del riesgo por la MAE y mejora continua.",
    "responsables": ["3.1 Oficial de Seguridad (CISO) / Dpto. TIC", "3.2 Dirección Nacional (MAE)"],
    "referencias": ["4.1 NORTIC A7:2016 (Capítulo V)", "4.2 NIST CSF v2.0 (Identify - Risk Assessment ID.RA)", "4.3 ISO/IEC 27005 (Guía)"],
    "politicas": [
        "5.1 Metodología de Riesgo: Se basa en cálculos cuantitativos y cualitativos, alineada al plan estratégico del organismo.",
        "5.2 Cálculo: Riesgo (R) = Impacto (I) x Probabilidad (P). Utilizando una escala del 1 al 4, generando riesgos de 1 a 16.",
        "5.3 Niveles de Riesgo: 1 a 6 (Bajo), 8 a 9 (Medio), 12 a 16 (Alto).",
        "5.4 Tratamiento: Todo riesgo debe someterse a acciones de mitigación, transferencia, anulación o aceptación.",
        "5.5 Aceptación: Todo riesgo residual requiere firma de la máxima autoridad confirmando que asume la responsabilidad operacional."
    ],
    "actividades": [
        ("CISO / Dpto. TIC", "6.1 Realiza el levantamiento periódico de activos y amenazas (virus, desastres naturales, fallas de hardware, robo)."),
        ("CISO / Dpto. TIC", "6.2 Aplica la fórmula (I x P) y elabora la Tabla de Apreciación de Riesgos."),
        ("CISO / Dpto. TIC", "6.3 Propone controles tecnológicos y administrativos (Firewalls, UPS, Políticas) para crear el Plan de Acción de Tratamiento."),
        ("Dirección Nacional", "6.4 Analiza el costo-beneficio de los controles propuestos y aprueba el presupuesto y el nivel de riesgo residual."),
        ("CISO / Dpto. TIC", "6.5 Monitorea los indicadores continuamente y presenta reportes anuales de la evolución de la matriz de riesgos.")
    ],
    "anexos": ["7.1 Tablas de Apreciación de Riesgos", "7.2 Plan de Tratamiento de Riesgos", "7.3 Acta de Aceptación de Riesgo Residual"],
    "tiempo": "Evaluación anual o al existir cambios estructurales."
}

operaciones_data = {
    "nombre": "Control de Operaciones e Inventario (Web, WiFi, Móviles, Cloud)",
    "codigo": "BNPHU-TIC-006",
    "preparado": "Departamento de TIC",
    "aprobado": "Dirección Nacional",
    "proposito": "Establecer las medidas de protección y control sobre las plataformas web, redes inalámbricas, dispositivos móviles corporativos, adopción de la nube y mantener un registro fiel de los activos de la institución.",
    "alcance_empieza": "Requerimiento de publicación web, despliegue de red WiFi, o ingreso de activos.",
    "alcance_incluye": "Actualizaciones de CMS (Seguridad Web), filtrado WiFi, cifrado y MDM para móviles, análisis de riesgos cloud y registro en el inventario oficial.",
    "alcance_termina": "Auditorías tecnológicas o descargo de los equipos.",
    "responsables": ["3.1 Dpto. TIC (Webmaster, Administrador de Redes)", "3.2 Dpto. de Activos Fijos"],
    "referencias": ["4.1 NORTIC A7:2016 (Capítulo VI)", "4.2 NIST CSF v2.0 (Identify - Asset Management ID.AM / Protect PR.PT)", "4.3 NORTIC A1", "4.4 NORTIC A2"],
    "politicas": [
        "5.1 Seguridad Web: Obligatorio el uso de HTTPS con certificados válidos (no autofirmados). Evaluaciones periódicas OWASP para el CMS y plugins.",
        "5.2 Servicios Wireless (WiFi): Prohibido usar protocolo WEP. El SSID institucional no debe revelar el nombre del organismo para redes ocultas. Se debe segmentar la red de visitantes de la red operativa.",
        "5.3 Dispositivos Móviles: Uso de sistemas MDM (Mobile Device Management) para borrado remoto. Prohibido compartir dispositivos asignados. Toda unidad extraviada se declara como incidente de seguridad.",
        "5.4 Servicios en la Nube: Ninguna información sensible nacional será almacenada en la nube si los servidores residen bajo leyes de estados ajenos sin análisis de riesgo y autorización de OGTIC.",
        "5.5 Inventario: Todos los activos físicos y de software deben llevar un registro y estar etiquetados."
    ],
    "actividades": [
        ("Webmaster", "6.1 Verifica suscripciones a listas de vulnerabilidades, aplica parches y elabora reportes mensuales de estado de aplicaciones web."),
        ("Administrador de Redes", "6.2 Configura las redes WiFi institucionales, aplicando filtros de acceso MAC (si aplica), IPS, y claves rotativas corporativas."),
        ("Dpto. TIC", "6.3 Inscribe dispositivos móviles en el software MDM antes de su entrega e instala el software de cifrado de almacenamiento."),
        ("Dpto. Activos / TIC", "6.4 Actualiza la base de datos de inventario con cada ingreso de hardware y software (SISTICGE/NORTIC A1).")
    ],
    "anexos": ["7.1 Matriz de Escaneo de Vulnerabilidades", "7.2 Listado de Activos Tecnológicos", "7.3 Acuerdos de Uso de Equipos (Móviles)"],
    "tiempo": "Monitoreo mensual de servicios. Actualización constante del inventario."
}

# --- GENERACION MASIVA ---
create_procedure("1-SASI/Procedimiento_SASI_BNPHU.docx", "portal-tic/public/Procedimiento_SASI_BNPHU.md", sasi_data)
create_procedure("2-Politicas_Seguridad/Administracion_Informacion.docx", "portal-tic/public/Administracion_Informacion.md", admin_info_data)
create_procedure("3-Controles_Acceso/Gestion_Accesos_y_Privilegios.docx", "portal-tic/public/Gestion_Accesos_y_Privilegios.md", acceso_data)
create_procedure("4-Continuidad_Negocio/Disponibilidad_y_Continuidad.docx", "portal-tic/public/Disponibilidad_y_Continuidad.md", continuidad_data)
create_procedure("5-Gestion_Riesgos/Gestion_Riesgos_Ciberseguridad.docx", "portal-tic/public/Gestion_Riesgos_Ciberseguridad.md", riesgos_data)
create_procedure("6-Control_Operaciones/Control_Operaciones_Activos.docx", "portal-tic/public/Control_Operaciones_Activos.md", operaciones_data)

print("Documentos DOCX y MD generados con exito y con contenido NORTIC/NIST completo!")
